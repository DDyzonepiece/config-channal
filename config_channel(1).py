import docx
import numpy as np
import pandas as pd
file_name_1=r'测试要求.docx'
file_name_2=r'引脚定义_01.docx'
file1=docx.Document(file_name_1)      #提取文档所有内容
file2=docx.Document(file_name_2)
print(file2.paragraphs[12].text)        ##？？？
table=file1.tables[1]
table_list = []
#读取测试要求的表格
for row in table.rows:  # 读每行
	row_content = []
	for cell in row.cells[1:6]:  # 读一行中的所有单元格   读第2 到第6列的元素 #修改，包含参数范围栏
		c = cell.text
		row_content.append(c)
	#print(row_content)
	table_list.append(row_content)
print(table_list[0][1])
print(len(file2.tables))
table2=file2.tables[0]
table_list2 = []
table_list3 = []
#读取引脚定义的表格
for row in table2.rows[1:]:  # 读每行
	  # 读一行中的第三个单元格
	c = row.cells[2].text
	d = row.cells[4].text
	table_list2.append(c)
	table_list3.append(d)


var_seq=[]
var_seq2=[]
#for item in table_list2:
#	temp=[]
#	temp=item.split('\n')
#	print(len(temp))
#	var_seq+=temp

for i in range(0,len(table_list2)):
	temp=[]
	temp=table_list2[i].split('\n')
	var_seq+=temp
	var_seq2+=[table_list3[i]]*len(temp)

print(var_seq)
print(len(var_seq))
#统计行数和配置参数
num_rows=len(table_list)
mm=[]
nn=[]
num_mn=[]
ran=[]
ran2=[]
#将统计的一些参数分别放入列表中
for i in range(1,num_rows):
	j=table_list[i][3][0]
	k=table_list[i][3][2]
	ran=table_list[i][4].split('~')
	ran[1]=int(ran[1])
	ran2.append(ran[1])
	#print(ran2)
	j=int(j)
	k=int(k)
	mm.append(j)
	nn.append(k)
	num_mn.append(j*k)     #num_mu 每个参数对应的分支点数量的列表

num_sum=sum(num_mn)#计算增加后的总行数
print(num_sum)
ran3=[]
ran4=[]
var_config=[]
#插空行填补
for i in range(num_rows-1):
	
	for j in range(mm[i]):
		for k in range(nn[i]):
			
			var_t=table_list[i+1][1]+'_%d'%(j+1)+'%d'%(k+1)
			var_config.append(var_t)
			ran3.append(ran2[i])
			if ran2[i]<=500:
				ran4.append('50psi')
			elif ran2[i]<=1000:
				ran4.append('100psi')
			elif ran2[i]<=2500:
				ran4.append('250psi')
			else:
				ran4.append('500psi')
			#print(ran4)
#print(var_config)
for i in range(len(ran4)):
	if var_seq2[i]!=None:
		ran4[i]=var_seq2[i]
		print(ran4[i])


#将空的行先配置到列表中


bb=[None]*5      ##???

n=2

for i in range(num_rows-1):
	
	if num_mn[i]>1:
		for j in range(num_mn[i]-1):
			table_list.insert(n,bb)          #在标签名称栏后面加上n*m-1行空格
	n+=num_mn[i]                           #n起始是2，后面每次加n*m

#生成dataframe结构数据类型
df1=pd.DataFrame(table_list[1:num_sum+1],columns=table_list[0])



#增加dataframe中的列
df1['通道配置参数']=var_config
df1['参数范围']=ran3
df1['传感器类型']=ran4
df1['编号']=var_seq
df1['偶型']=var_seq2
df1.to_excel(r'config.xlsx', index=None, columns=None)
#print(df1)
