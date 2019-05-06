import docx
import numpy as np
import pandas as pd
import openpyxl
import datetime
from openpyxl.utils import get_column_letter
from openpyxl.styles import Color, PatternFill, Font, Border
from openpyxl.styles import colors


"""
Info:
	将输入的最大压力变为对应PSI
Args:
    p_input-输入的最大压力
Returns:
    各个对应的PSI
Modify:
    2019-4-18
"""
def psi_out(p_input):
	psi=(int(p_input)-101.25)/6.895
	#print(psi)
	if psi<5:
		return 5
	elif 5<=psi<15:
		return 15
	elif 15<=psi<30:
		return 30
	elif 30<=psi<100:
		return 100
	elif 100<=psi<250:
		return 250
	elif 250<=psi<500:
		return 500
	elif 500 <= psi < 750:
		return 750
	else:
		return 0
"""
Info: 
	统计出合并单元格的范围,可以选定第几列
Args:
    worksheet-需要输入的表名
    column-需要检查的合并的列数，没有则检查所有列
Returns:
    merged_list-合并的单元格的坐标元组
Modify:
    2019-4-18
"""
def get_merged_range(worksheet,column=None):
	m_list = worksheet.merged_cells

	# 合并单元格的起始行坐标、终止行坐标
	merged_list=[]
	for m_area in m_list:
		temp_tuple = (m_area.min_row, m_area.max_row, m_area.min_col, m_area.max_col)
		merged_list.append(temp_tuple)

	if column!=None:
		temp_list=[]
		for item in merged_list:

			if item[2]==column and item[3]==column:
				temp_list.append(item)
		merged_list=temp_list

	return  merged_list

"""
Info: 
	统计所选的那一列所有空的单元格坐标
Args:
    worksheet-需要输入的表名
    column-需要检查空的列数
Returns:
    empty_list-空的单元格坐标，认定None和只有一个空格都是空
Modify:
    2019-4-18
"""
def get_empty_list(worksheet,column):
	num_r = worksheet.max_row
	empty_list = []
	for i in range(1, num_r + 1):
		temp = worksheet.cell(row=i, column=column).value
		if temp == None or temp == ' ':
			empty_list.append(i)

	return empty_list




file_name_1=r'测试要求.docx'
file_name_2=r'引脚定义_01.docx'
file1=docx.Document(file_name_1)
file2=docx.Document(file_name_2)
#print(file2.paragraphs[12].text)
#用字典类似于json形式去装配置好的表格
#excel是输出的表现形式，内存里用字典来表达数据关系
table_dir_1={}

table=file1.tables[1]
table_list_1 = []
range_list=[]
#读取测试要求的表格
for row in table.rows:  # 读每行
	row_content = []
	range_list.append(row.cells[5].text)
	for cell in row.cells[1:7]:  # 读一行中的所有单元格
		c = cell.text
		row_content.append(c)
	# 用二维列表去表示读取的表格
	table_list_1.append(row_content)
# print(table_list_1[2][0])
# print(len(file2.tables))
table2=file2.tables[0]
table_list_2 = []
#读取引脚定义的表格
for row in table2.rows[1:]:  # 读每行
	  # 读一行中的第三个单元格
	c = row.cells[2].text

	table_list_2.append(c)


var_seq=[]
for item in table_list_2:
	temp=[]
	temp=item.split('\n')
	var_seq+=temp

#print(var_seq)
#print(len(var_seq))
#统计行数和配置参数
num_rows=len(table_list_1)
mm=[]
nn=[]
num_mn=[]
#之后要改成split分割形式不然通道是2位数就没法用
#将统计的一些参数分别放入列表中
for i in range(1,num_rows):
	j=table_list_1[i][3][0]
	k=table_list_1[i][3][2]
	j=int(j)
	k=int(k)
	mm.append(j)
	nn.append(k)
	num_mn.append(j*k)

num_sum=sum(num_mn)#计算增加后的总行数
print(num_sum)
#var_config是用来存放配置好的变量名列表
var_config=[]
#var_config_2dim是以2维列表的方式来存配置好变量名
var_config_2dim=[]
#插空行填补
for i in range(num_rows-1):
	temp_list=[]
	for j in range(mm[i]):
		for k in range(nn[i]):
			
			var_t=table_list_1[i+1][1]+'_%d'%(j+1)+'%d'%(k+1)
			temp_list.append(var_t)
			var_config.append(var_t)
	var_config_2dim.append(temp_list)

#table_list_1[0].append('量程（表压）')
max_range_list=[]
#分割范围
for item in range_list[1:]:
	temp=[]
	temp=item.split('~')
	#print(temp)
	max_range_list.append(temp[1])

#print(max_range_list)

#psi_list

psi_list=[psi_out(x) for x in max_range_list]
#print(psi_list)
#将非压力的参数改为空格
for i,item in enumerate(psi_list):
	if table_list_1[i+1][1][0]!='P':
		psi_list[i]=' '
#print(psi_list)
value_columns=[table_list_1[0][0],
			   table_list_1[0][2],
			   table_list_1[0][3],
			   table_list_1[0][4],
			   table_list_1[0][5],
			   '量程（表压）',
			   '通道配置参数']

for i,item in enumerate(table_list_1[1:]):

	table_dir_1[item[1]]={table_list_1[0][0]:item[0], \
						  table_list_1[0][2]: item[2],\
						  table_list_1[0][3]: item[3], \
						  table_list_1[0][4]: item[4], \
						  table_list_1[0][5]: item[5], \
						  '量程（表压）':psi_list[i],\
						  '通道配置参数':var_config_2dim[i],\
						  }

#print(table_dir_1)
# psi_list_extend=[]
# for i in range(len(psi_list)):
# 	item=[psi_list[i]]*num_mn[i]
# 	psi_list_extend.extend(item)
#
# print(psi_list_extend)
# print(len(psi_list_extend))


#将空的行先配置到列表中
bb=[None]*6

n=2

for i in range(num_rows-1):
	
	if num_mn[i]>1:
		for j in range(num_mn[i]-1):
			table_list_1.insert(n,bb)
			psi_list.insert(n-1,None)
	n+=num_mn[i]

#生成dataframe结构数据类型
df1=pd.DataFrame(table_list_1[1:num_sum+1],columns=table_list_1[0])


#print(psi_list)

#增加dataframe中的列
df1['量程（表压）（psi）']=psi_list
df1['通道配置参数']=var_config
df1['编号']=var_seq

df1.to_excel(r'config.xlsx', index=None, columns=None)
#print(df1)


#开始处理excel并从中提取信息
excel_file_path=r"2007a_test.xlsx"
workbook = openpyxl.load_workbook(excel_file_path)
# workbook = openpyxl.Workbook(path)
name_list = workbook.sheetnames
sheet_index=0
worksheet = workbook[name_list[sheet_index]]
num_r=worksheet.max_row
#print(num_r)
#print(worksheet[10].value)

merged_list=get_merged_range(worksheet,column=10)
empty_list=get_empty_list(worksheet,10)
temp_list = empty_list.copy()#这个非常重要否则迭代会出现问题
#print(len(empty_list))
for m_item in merged_list:
	#print(m_item)
	for e_item in temp_list:
		if e_item>=m_item[0] and e_item<=m_item[1]:
			empty_list.remove(e_item)



#newlist=[]
# for m_item in merged_list:
# 	newlist += [ i for i in empty_list if i>= m_item[0]and i<=m_item[1]]
#
# list1 = [i for i in empty_list if i not in newlist]

#print(len(empty_list))
#print(empty_list)

"""
Info: 
	获取提供的配置表中各个PSI大小的位子的字典
Args:
    worksheet-需要输入的表名
    column-需要获取的位子列数，默认是3
Returns:
    psi_dir-包含各个psi位子的字典
Modify:
    2019-4-19
"""
def get_psi_dir(worksheet,empty_list,column=3):
	psi_dir={}
	psi_5_list=[]
	psi_15_list=[]
	psi_30_list=[]
	psi_100_list=[]
	psi_250_list=[]
	psi_500_list=[]
	psi_750_list=[]
	for item in empty_list:
		temp = worksheet.cell(row=item, column=column).value
		if temp=='± 5 psi':
			psi_5_list.append(item)
		elif temp=='± 15 psi':
			psi_15_list.append(item)
		elif temp=='± 30 psi':
			psi_30_list.append(item)
		elif temp=='100 psi':
			psi_100_list.append(item)
		elif temp=='250 psi':
			psi_250_list.append(item)
		elif temp=='500 psi':
			psi_500_list.append(item)
		elif temp=='750 psi':
			psi_750_list.append(item)
	psi_dir['psi_5_list']=psi_5_list
	psi_dir['psi_15_list'] = psi_15_list
	psi_dir['psi_30_list'] = psi_30_list
	psi_dir['psi_100_list'] = psi_100_list
	psi_dir['psi_250_list'] = psi_250_list
	psi_dir['psi_500_list'] = psi_500_list
	psi_dir['psi_750_list'] = psi_750_list

	return psi_dir


psi_dir=get_psi_dir(worksheet,empty_list,column=3)


psi_kinds=[5,15,30,100,250,500,750]
# curr_time=datetime.datetime.now()
# curr_time=curr_time.strftime('%Y-%m-%d %H:%M:%S')
# text=curr_time+'\n'
# text+='配置表中压力测点数量统计:\n'
# for i,value in enumerate(psi_dir.values()):
# 	text += "%d psi:%d \n" % (psi_kinds[i], len(value))



"""
Info: 
	统计需求定义表中和提供通道的表格的统计信息
Args:
    psi_dir-是用来装各个psi位置列表的字典
    table_dir_1-主要用来装从表格中获得的和配好变量后的数据字典
    psi_kinds-是各个psi值得列表
Returns:
    psi_supply_count-配置表的psi提供的统计数据，是各个对应的psi的数量列表
    psi_demand_count-测试要求中对psi需求的统计数据，是各个对应的psi的数量列表
    				 该列表每个对应的psi，有两个数据组成一个元组，第一个数是所有
    				 该psi需求数量，第二个数是需求数量中精度要求小于0.35%
Modify:
    2019-4-28
"""
def psi_count_show(psi_dir,table_dir_1,psi_kinds):
	#psi_kinds = [5, 15, 30, 100, 250, 500, 750]
	psi_supply_count=[]
	psi_demand_count=[]
	for i, value in enumerate(psi_dir.values()):
		psi_supply_count.append(len(value))
		#text += "%d psi:%d \n" % (psi_kinds[i], len(value))

	for item in psi_kinds:
		#temp_list=[]
		count_all=0
		count_acc=0
		for g_value in table_dir_1.values():

			if g_value['量程（表压）']==item:
				count_all+=len(g_value['通道配置参数'])

				try:

					acc=g_value['总精度要求']

					acc=acc[1:-1]
					#print(acc)


					acc=float(acc)


					if acc<=0.35:
						count_acc+=len(g_value['通道配置参数'])
				except Exception as e:

					print("压力精度设置有问题，请检查压力精度设置表格")

		psi_demand_count.append((count_all,count_acc))

	return psi_supply_count,psi_demand_count

psi_supply_count,psi_demand_count=psi_count_show(psi_dir,table_dir_1,psi_kinds)

#print(psi_demand_count)
#print(table_dir_1)
curr_time=datetime.datetime.now()
curr_time=curr_time.strftime('%Y-%m-%d %H:%M:%S')
text=curr_time+'\n'
print('配置表中压力测点数量统计:')
print(psi_supply_count)
print('测试要求中压力测点数量统计(所有数量，不可拓展数量):')
print(psi_demand_count)



"""
Info: 
	判断是否需要拓展和拓展后通道数是否足够
Args:
    psi_supply_count-配置表的psi提供的统计数据，是各个对应的psi的数量列表
    psi_demand_count-测试要求中对psi需求的统计数据，是各个对应的psi的数量列表
    				 该列表每个对应的psi，有两个数据组成一个元组，第一个数是所有
    				 该psi需求数量，第二个数是需求数量中精度要求小于0.35%
    psi_kinds-是各个psi值得列表
    text-记录的变量
Returns:
	no_expand_psi-布尔型如果为True则不需要扩展可配置
	expand_psi_enough-布尔型如果为True则采用拓展的方式可以配置
	extra_number-用来装对应的各个psi需要拓展到下一层的数量
	text-记录的变量
Modify:
    2019-4-29
"""
def config_check(psi_supply_count,psi_demand_count,psi_kinds,text):

	no_expand_psi=True
	expand_psi_enough=True
	extra_number = [0] * len(psi_kinds)  # 用来存放每次多出来的需要拓展的数量
	text += "按照不扩大量程方式配置，检查配置表中数量是否都满足配置条件： \n"
	for i,item in enumerate(psi_demand_count):

		if item[0]<=psi_supply_count[i]:
			text+="%d psi 配置表中数量满足配置条件 \n"%(psi_kinds[i])

		else:
			text += "%d psi 配置表中数量不满足配置条件 \n" % (psi_kinds[i])
			no_expand_psi=False

	if no_expand_psi==False:


		while True:
			whether_expand = input('按照对应量程以无法满足，是否按照扩大量程方式配置?  y/n \n')
			if whether_expand=='n':

				print('选择不扩大量程配置，提供的配置通道表不满足条件，请增加对应缺少的通道数')
				break
			elif whether_expand=='y':
				text +='已选择扩大量程方式配置，只有精度大于0.35%的变量才会使用扩大量程\n'
				print('已选择扩大量程方式配置，只有精度大于0.35%的变量才会使用扩大量程')
				break
			else:
				print('输入错误，请重新输入！')




		for  i,item in enumerate(psi_demand_count):

			if  i==0:
				if item[0]<=psi_supply_count[i]:
					text += "%d psi 配置表中对应量程数量满足配置条件,不需要拓展 \n" % (psi_kinds[i])



				elif item[1] < psi_supply_count[i] :
					text += "%d psi 配置表中数量能满足测试要求中精度小于0.35%的数量，但需要%d个扩大量程 \n"\
							% (psi_kinds[i],item[1]- psi_supply_count[i])
					extra_number[i]=item[1]- psi_supply_count[i]

				else:

					text += "%d psi 配置表中精数量小于测试要求中精度小于0.35%的数量，少%d个 \n" \
							% (psi_kinds[i], item[1] - psi_supply_count[i])
					extra_number[i] = item[0] - item[1]
					expand_psi_enough = False


			if i>0:

				if item[0]+extra_number[i-1]<= psi_supply_count[i] :

					text+="%d psi 配置表中数量满足拓展后配置条件 \n"%(psi_kinds[i])

				elif item[1]+extra_number[i-1] <=psi_supply_count[i] :

					extra_number[i]=item[0]+extra_number[i-1]-psi_supply_count[i]
					text += "%d psi 配置表中数量满足拓展后配置条件，但需要有%d个去拓展到更大量程\n" \
							% (psi_kinds[i],item[0]+extra_number[i-1]-psi_supply_count[i])

				else :

					extra_number[i] =item[0]-item[1]
					text += "%d psi 配置表中数量不满足拓展后配置条件，需要有增加%d个通道\n" \
							% (psi_kinds[i], item[1] + extra_number[i-1] - psi_supply_count[i])
					expand_psi_enough = False
	return no_expand_psi,expand_psi_enough,extra_number,text


no_expand_psi,expand_psi_enough,extra_number,text=config_check(psi_supply_count,psi_demand_count,psi_kinds,text)

print(extra_number)
#print(table_dir_1)
if no_expand_psi:

	temp_i=0
	changed_list=[]
	for psi_key,psi_value in psi_dir.items():
		var_psi=[]#用来存放需要填入表格的变量列表
		#第一个循环将需要配置的压力参数全部提出来
		for key,value in table_dir_1.items():
			is_pressure=value['量程（表压）']

			if is_pressure==psi_kinds[temp_i]:
				var_psi.extend(value['通道配置参数'])


		temp_i+=1

		if len(var_psi)<=len(psi_dir[psi_key]) and len(var_psi)!=0:

			#这个循环就是将需要填充的变量全部放进表格中，按psi匹配
			for i,item in enumerate(var_psi):
				#exl=get_column_letter(10)+str(psi_dir['psi_30_list'][i])
				worksheet.cell(row=psi_dir[psi_key][i], column=10).value=item
				changed_list.append(psi_dir[psi_key][i])

elif expand_psi_enough:
	temp_i = 0
	extra_number.insert(0,0)#在记录每个需要拓展的个数的列表最前端加个0便于循环

	print(extra_number)
	changed_list = []#用来记录这次配置的通道的位置
	var_add_list=[]#用来存放上一层扩展后放在本层的变量
	var_remove_list=[]#用来存放本层需要剔除的，需要拓展到下一层的变量
	for psi_key, psi_value in psi_dir.items():
		var_psi = []  # 用来存放需要填入表格的变量列表
		#第一次只用判断本层就可以了，之后需要判断上一层是否有拓展和本层是否有拓展
		if extra_number[temp_i]==0 and extra_number[temp_i+1]==0:

			# 第一个循环将需要配置的压力参数全部提出来
			for key, value in table_dir_1.items():
				is_pressure = value['量程（表压）']
				if is_pressure == psi_kinds[temp_i]:
					var_psi.extend(value['通道配置参数'])

			var_add_list = []
			temp_i += 1

			if len(var_psi) <= len(psi_dir[psi_key]) and len(var_psi) != 0:#这句其实并不必要，因为之前已经做了配置检查

				# 这个循环就是将需要填充的变量全部放进表格中，按psi匹配
				for i, item in enumerate(var_psi):
					# exl=get_column_letter(10)+str(psi_dir['psi_30_list'][i])
					worksheet.cell(row=psi_dir[psi_key][i], column=10).value = item
					changed_list.append(psi_dir[psi_key][i])

		else:


			var_psi_35= []#提取出精度小于0.35%的变量名，优先级最小
			var_psi_no35=[]#提取出精度大于0.35%的变量名，优先级最差

			# 第一个循环将需要配置的压力参数全部提出来
			for key, value in table_dir_1.items():
				is_pressure = value['量程（表压）']
				if is_pressure == psi_kinds[temp_i]:

					acc = value['总精度要求']

					acc = acc[1:-1]

					acc = float(acc)

					if acc <= 0.35:
						var_psi_35.extend(value['通道配置参数'])
						#print('通道配置参数')

					else:
						var_psi_no35.extend(value['通道配置参数'])
			# print('通道配置参数')
			# print(len(var_psi_35))
			# print(len(var_psi_no35))
			if extra_number[temp_i+1]!=0:#为本层有需要移除的
				#print(temp_i+1)
				var_remove_list=var_psi_no35[len(var_psi_no35)-extra_number[temp_i+1]:]
				var_psi_no35=var_psi_no35[:len(var_psi_no35)-extra_number[temp_i+1]]
				if var_add_list!=[]:
					print('%d psi本层额外增加变量'%(psi_kinds[temp_i]))
					print(var_add_list)
				var_psi.extend(var_add_list)#先将上一层拓展的放进变量列表
				var_psi.extend(var_psi_35)#再将本层小于0.35%放进去
				var_psi.extend(var_psi_no35)#最后将本层大于0.35%的没移除的放进去
				var_add_list = var_remove_list  # 为上次超出的需要拓展变量名，优先级最大
				# print('移除的数量')
				# print(len(var_remove_list))
				# print(len(var_psi_no35))
			else:
				var_psi.extend(var_add_list)  # 先将上一层拓展的放进变量列表
				print('%d psi本层额外增加变量'%(psi_kinds[temp_i]))
				print(var_add_list)
				var_psi.extend(var_psi_35)  # 再将本层小于0.35%放进去
				var_psi.extend(var_psi_no35)  # 最后将本层大于0.35%的没移除的放进去
				var_add_list =[]#因为本层没有需要拓展的所以将剔除的给更新

			temp_i += 1
			if len(var_psi) <= len(psi_dir[psi_key]) and len(var_psi) != 0:#这句其实并不必要，因为之前已经做了配置检查

				# 这个循环就是将需要填充的变量全部放进表格中，按psi匹配
				for i, item in enumerate(var_psi):
					# exl=get_column_letter(10)+str(psi_dir['psi_30_list'][i])
					worksheet.cell(row=psi_dir[psi_key][i], column=10).value = item
					changed_list.append(psi_dir[psi_key][i])



else:
	print('请先修改配置表格满足拓展要求')












config_list=[]
for item in changed_list:
	worksheet.cell(row=item, column=12).value = worksheet.cell(row=item, column=9).value\
												+'_'+worksheet.cell(row=item, column=10).value

	config_list.append((worksheet.cell(row=item, column=10).value,
						worksheet.cell(row=item, column=9).value,
						worksheet.cell(row=item, column=8).value))

workbook.save('test.xlsx')#将已经配好的表格保存下来
"""
Info: 
	将给的配置表格中的匹配好的提出放到自己表格中
Args:
    config_list-包含'PINOUT'，'Channel Number'和'测点名称'的元组列表
Returns:
    保存一个config_channel.xlsx文件
Modify:
    2019-4-19
"""
def config_channel(config_list):
	excel_file_path=r"merge.xlsx"
	workbook = openpyxl.load_workbook(excel_file_path)
	name_list = workbook.sheetnames
	sheet_index = 0
	worksheet = workbook[name_list[sheet_index]]
	worksheet['J1']='PINOUT'
	worksheet['K1'] = 'Channel Number'
	num_r = worksheet.max_row

	for i in range(1,num_r+1):


		for item in config_list:
			#print(item)
			if worksheet.cell(row=i, column=8).value==item[0]:
				worksheet.cell(row=i, column=10).value=item[1]
				worksheet.cell(row=i, column=11).value = item[2]
				config_list.remove(item)
				break

	workbook.save('config_channel.xlsx')


config_channel(config_list)
#print(len(empty_list))
#workbook.save('test.xlsx')
print(text)